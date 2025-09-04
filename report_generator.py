#!/usr/bin/env python3
"""
Temperature report generator with concurrent GUI progress per-vehicle.
Saves per-vehicle DOCX reports named temp_report_<vehicle>.docx inside each vehicle folder.
"""

import os
import re
import time
import traceback
import threading
from collections import Counter
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# =========================
# Configuration
# =========================
MAX_CONCURRENT_VEHICLES = 5  # Change this number to control concurrency

# Pull the download root from Script 2 so we never hardcode paths
try:
    import email_reader_attachment_download as script2
    DEFAULT_DOWNLOAD_ROOT = getattr(
        script2, "ROOT_DOWNLOAD_DIR",
        os.path.join(os.path.dirname(__file__), "download")
    )
except Exception:
    DEFAULT_DOWNLOAD_ROOT = os.path.join(os.path.dirname(__file__), "download")

# =========================
# Tkinter progress UI
# =========================
import tkinter as tk
from tkinter import ttk

class ReportProgressGUI:
    def __init__(self, vehicle_names, auto_close_secs=2):
        self.root = tk.Tk()
        self.root.title("Temperature Report Generator")
        self.root.geometry("760x520")
        self.auto_close_secs = auto_close_secs

        tk.Label(self.root, text="Vehicle Report Progress", font=("Arial", 14, "bold")).pack(pady=(8, 0))

        container = tk.Frame(self.root)
        container.pack(fill="both", expand=True, padx=8, pady=(6, 2))

        self.canvas = tk.Canvas(container, borderwidth=0)
        self.scrollbar = ttk.Scrollbar(container, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.vehicle_widgets = {}
        for name in vehicle_names:
            row = tk.Frame(self.scrollable_frame, pady=6)
            row.pack(fill="x", padx=6)

            lbl_name = tk.Label(row, text=name, font=("Arial", 10, "bold"), width=22, anchor="w")
            lbl_name.pack(side="left", padx=(2, 6))

            pb = ttk.Progressbar(row, orient="horizontal", length=300, mode="determinate")
            pb.pack(side="left", padx=(0, 8))

            lbl_status = tk.Label(row, text="Waiting…", font=("Arial", 9), anchor="w", wraplength=300, justify="left")
            lbl_status.pack(side="left", fill="x", expand=True)

            self.vehicle_widgets[name] = {"progress": pb, "status": lbl_status, "name_label": lbl_name}

        footer = tk.Frame(self.root)
        footer.pack(fill="x", pady=(4, 8))

        self.global_status = tk.Label(footer, text="Ready", font=("Arial", 11))
        self.global_status.pack(side="left", padx=8)

        self.close_btn = tk.Button(footer, text="Close Now", command=self._close_now, state="disabled")
        self.close_btn.pack(side="right", padx=8)

        self._closed = False

    def start(self):
        try:
            self.root.mainloop()
        except Exception:
            pass

    def _close_now(self):
        self._closed = True
        try:
            self.root.destroy()
        except Exception:
            pass

    def update_progress(self, vehicle_name: str, percent: int, status: str):
        def _apply():
            if vehicle_name in self.vehicle_widgets:
                w = self.vehicle_widgets[vehicle_name]
                w["progress"]["value"] = max(0, min(100, percent))
                w["status"].config(text=status)
                self.global_status.config(text=f"Processing: {vehicle_name} — {status}")
        try:
            self.root.after(0, _apply)
        except Exception:
            pass

    def mark_vehicle_done(self, vehicle_name: str, msg="Done ✅"):
        def _apply():
            if vehicle_name in self.vehicle_widgets:
                w = self.vehicle_widgets[vehicle_name]
                w["progress"]["value"] = 100
                w["status"].config(text=msg)
                self.global_status.config(text=f"Completed: {vehicle_name}")
        try:
            self.root.after(0, _apply)
        except Exception:
            pass

    def mark_done(self, msg="✅ All reports generated"):
        def _apply():
            self.global_status.config(text=msg)
            self.close_btn.config(state="normal")
            if not self._closed and self.auto_close_secs is not None and self.auto_close_secs >= 0:
                self.root.after(int(self.auto_close_secs * 1000), self._close_now)
        try:
            self.root.after(0, _apply)
        except Exception:
            pass

# =========================
# Helpers
# =========================
def get_closest_column(df_cols, candidates):
    for name in candidates:
        if name in df_cols:
            return name
    return None

def set_page_border(section):
    sectPr = section._sectPr
    for elem in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(elem)
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

def _safe_basename_no_ext(path: str) -> str:
    base = os.path.basename(path)
    stem, _ = os.path.splitext(base)
    return stem

# =========================
# Core per-vehicle generator
# =========================
def generate_report_for_vehicle(vehicle_folder: str, progress_cb=None):
    vehicle_name = os.path.basename(vehicle_folder)
    # Delete old DOCX
    for f in os.listdir(vehicle_folder):
        if f.lower().endswith(".docx") and f.startswith("temp_report_"):
            try:
                os.remove(os.path.join(vehicle_folder, f))
            except Exception:
                pass

    xlsx_files = [os.path.join(vehicle_folder, f) for f in os.listdir(vehicle_folder) if f.lower().endswith(".xlsx")]
    if not xlsx_files:
        if progress_cb:
            progress_cb(vehicle_name, 0, f"❌ No Excel files found in {vehicle_folder}")
        return

    try:
        all_dfs = []
        total_rows = 0
        for f in xlsx_files:
            try:
                df_tmp = pd.read_excel(f)
            except Exception:
                df_tmp = pd.DataFrame()
            if df_tmp.empty:
                continue
            total_rows += len(df_tmp)
            all_dfs.append((f, df_tmp))

        if total_rows == 0:
            if progress_cb:
                progress_cb(vehicle_name, 0, "❌ All Excel files are empty")
            return

        if progress_cb:
            progress_cb(vehicle_name, 1, f"Found {len(all_dfs)} file(s), {total_rows} rows…")

        doc = Document()
        set_page_border(doc.sections[0])
        overall_max_val = None
        overall_max_count = None
        overall_max_date = None
        processed_rows = 0

        for file_path, df in all_dfs:
            if progress_cb:
                pct_now = int((processed_rows / total_rows) * 100) if total_rows else 0
                progress_cb(vehicle_name, pct_now, f"Reading {os.path.basename(file_path)} …")

            match = re.search(r'Parsed_(\d+)', os.path.basename(file_path))
            report_id = match.group(1) if match else _safe_basename_no_ext(file_path)

            temp_columns = [c for c in df.columns if re.search(r'battery.*temp.*\d+', c, flags=re.IGNORECASE)]
            if not temp_columns:
                continue

            soc_col = get_closest_column(df.columns, ['batteryStateOfCharge', 'battery_state_of_charge', 'SoC'])
            if not soc_col or 'createdAt' not in df.columns:
                continue

            df[temp_columns] = df[temp_columns].apply(pd.to_numeric, errors='coerce').clip(-50, 300)
            df[soc_col] = pd.to_numeric(df[soc_col], errors='coerce')
            df['createdAt'] = pd.to_datetime(df['createdAt'], errors='coerce')
            df.dropna(subset=[soc_col, 'createdAt'], inplace=True)
            if df.empty:
                continue

            df['MaxTemp'] = df[temp_columns].max(axis=1)
            df['MinTemp'] = df[temp_columns].min(axis=1)
            df['TempImbalance'] = df['MaxTemp'] - df['MinTemp']
            try:
                df['MaxTempCell'] = df[temp_columns].idxmax(axis=1)
                df['MinTempCell'] = df[temp_columns].idxmin(axis=1)
            except Exception:
                df['MaxTempCell'] = None
                df['MinTempCell'] = None

            df['date'] = df['createdAt'].dt.date
            if df['date'].dropna().empty:
                continue
            try:
                best_date = df['date'].value_counts().idxmax()
            except Exception:
                continue
            group = df[df['date'] == best_date].copy()
            if group.empty:
                continue

            # Metrics
            start_soc = group[soc_col].iloc[0]
            end_soc = group[soc_col].iloc[-1]
            most_max_cell = Counter(group['MaxTempCell']).most_common(1)[0][0] if not group['MaxTempCell'].isnull().all() else None
            most_min_cell = Counter(group['MinTempCell']).most_common(1)[0][0] if not group['MinTempCell'].isnull().all() else None
            max_imbalance = group['TempImbalance'].max()
            imbalance_count = int((group['TempImbalance'] == max_imbalance).sum())
            if (overall_max_val is None) or (max_imbalance > overall_max_val):
                overall_max_val = float(max_imbalance)
                overall_max_count = imbalance_count
                overall_max_date = best_date

            # Heading
            h = doc.add_heading(f"TEMPERATURE PROFILE FOR {report_id}", level=0)
            h.runs[0].font.size = Pt(14)
            h.runs[0].bold = True
            doc.add_heading(f"Date: {best_date}", level=1)

            # Plot
            x = range(len(group))
            soc_vals = group[soc_col].values
            y_max = group['MaxTemp'].values
            y_min = group['MinTemp'].values
            y_imb = group['TempImbalance'].values

            fig, ax1 = plt.subplots(figsize=(8, 4.5))
            ax1.plot(x, y_max, color="blue", linewidth=1.2, label="MaxTemp (°C)")
            ax1.plot(x, y_min, color="green", linewidth=1.2, label="MinTemp (°C)")
            ax1.set_xlabel("BatteryStateOfCharge (SoC)")
            ax1.set_ylabel("Temperature (°C)")
            try:
                ax1.set_xticks(x)
                ax1.set_xticklabels([f"{v:.0f}" for v in soc_vals], rotation=45, ha="right")
                if len(soc_vals) > 25:
                    step = max(1, len(soc_vals) // 25)
                    for i, label in enumerate(ax1.xaxis.get_ticklabels()):
                        if i % step != 0:
                            label.set_visible(False)
            except Exception:
                pass
            ax2 = ax1.twinx()
            ax2.plot(x, y_imb, color="red", linewidth=1.2, label="TempImbalance (°C)")
            ax2.set_ylabel("Imbalance (°C)")
            lines_1, labels_1 = ax1.get_legend_handles_labels()
            lines_2, labels_2 = ax2.get_legend_handles_labels()
            ax1.legend(lines_1 + lines_2, labels_1 + labels_2, loc="upper left", frameon=False)
            plt.title(f"Battery Temperatures & Imbalance — {best_date}", fontsize=11, weight="bold")
            plt.grid(True, linestyle="--", linewidth=0.5, alpha=0.7)
            plt.tight_layout()
            img_stream = BytesIO()
            plt.savefig(img_stream, bbox_inches="tight", dpi=120)
            plt.close(fig)
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(6.5))

            # Table
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light List Accent 1'
            table.cell(0, 0).text = "Start BatteryStateOfCharge"
            table.cell(0, 1).text = f"{start_soc:.2f}"
            table.cell(1, 0).text = "End BatteryStateOfCharge"
            table.cell(1, 1).text = f"{end_soc:.2f}"
            table.cell(2, 0).text = "Most data point of MaxTempCell"
            table.cell(2, 1).text = str(most_max_cell)
            table.cell(3, 0).text = "Most data point of MinTempCell"
            table.cell(3, 1).text = str(most_min_cell)
            cell0 = table.cell(4, 0)
            cell1 = table.cell(4, 1)
            cell0.text = "Max imbalance logged"
            cell1.text = f"{max_imbalance:.2f} (Count: {imbalance_count})"
            for cell in (cell0, cell1):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), "FF0000")
                tcPr.append(shd)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            table.cell(5, 0).text = "Total Records (this date)"
            table.cell(5, 1).text = str(len(group))

            doc.add_page_break()
            set_page_border(doc.sections[-1])

            processed_rows += len(group)
            if total_rows > 0 and progress_cb:
                pct = int((processed_rows / total_rows) * 100)
                progress_cb(vehicle_name, pct, f"Processing {_safe_basename_no_ext(file_path)} … {pct}%")

            time.sleep(0.02)

        if overall_max_val is not None:
            p = doc.paragraphs[0].insert_paragraph_before()
            run = p.add_run(
                f"MAX imbalance observed across all Days: {overall_max_val:.2f} "
                f"(Count: {overall_max_count}), on Date: {overall_max_date}"
            )
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)

        out_path = os.path.join(vehicle_folder, f"temp_report_{os.path.basename(vehicle_folder)}.docx")
        doc.save(out_path)
        if progress_cb:
            progress_cb(vehicle_name, 100, f"✅ Saved {out_path}")

    except Exception as e:
        if progress_cb:
            progress_cb(vehicle_name, 0, f"❌ Error: {e}")
        traceback.print_exc()


# =========================
# Batch driver (multi-threaded)
# =========================

# Running lock / flag to avoid concurrent runs
_running_lock = threading.Lock()
_running = False

def generate_all_reports(download_root: str | None = None, show_gui: bool = True, max_workers: int = MAX_CONCURRENT_VEHICLES):
    global _running
    # Prevent re-entrance
    with _running_lock:
        if _running:
            print("⚠️ generate_all_reports already running — skipping new call.")
            return
        _running = True

    try:
        root = download_root or DEFAULT_DOWNLOAD_ROOT

        if not os.path.isdir(root):
            print(f"❌ Download folder not found: {root}")
            return

        vehicle_folders = [
            os.path.join(root, d)
            for d in os.listdir(root)
            if os.path.isdir(os.path.join(root, d))
        ]
        if not vehicle_folders:
            print("❌ No vehicle folders found")
            return

        vehicle_names = [os.path.basename(p) for p in vehicle_folders]
        gui = ReportProgressGUI(vehicle_names) if show_gui else None

        def progress_wrapper(v_name, pct, msg):
            if gui:
                gui.update_progress(v_name, pct, msg)
            else:
                print(f"[{v_name}] {pct}% - {msg}")

        def task(folder):
            name = os.path.basename(folder)
            progress_wrapper(name, 0, f"📂 Starting {name}")
            generate_report_for_vehicle(folder, progress_cb=progress_wrapper)
            if gui:
                gui.mark_vehicle_done(name, "Done ✅")
            return name

        def run_executor():
            try:
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(task, f): f for f in vehicle_folders}
                    for fut in as_completed(futures):
                        try:
                            fut.result()
                        except Exception as e:
                            print("❌ Error processing vehicle:", e)
            except Exception as e:
                print("❌ Executor error:", e)
            finally:
                # final GUI update happens in main thread via mark_done
               

                if gui:
                    gui.mark_done("✅ All reports generated")
                else:
                    print("✅ All reports generated")

        # Start worker thread (non-daemon) and join it to keep lifecycle clear
        t = threading.Thread(target=run_executor, daemon=False)
        t.start()

        # If GUI enabled: start it (this blocks until GUI closed). After GUI closes, join worker thread.
        if gui:
            try:
                gui.start()
            except Exception:
                pass
            # GUI closed — wait for worker to finish
            t.join()
        else:
            # No GUI: just wait for worker to finish
            t.join()

    finally:
        with _running_lock:
            _running = False


# =========================
# Standalone
# =========================
if __name__ == "__main__":
    generate_all_reports(show_gui=True)
